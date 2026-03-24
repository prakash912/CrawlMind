"""
URL discovery (accurate) → crawl one by one → one DOCX per URL (named by URL) → OpenAI format.
"""
import asyncio
import aiohttp          # async HTTP client for fetching pages
import hashlib          # for short hashes in filenames when URL is long or has query
import json
import os
import re               # regex for cleaning URLs and markdown
import uuid
from pathlib import Path
from datetime import datetime
from io import BytesIO  # in-memory buffer for building DOCX before writing to disk
from urllib.parse import urljoin, urlparse, urlunparse

from bs4 import BeautifulSoup  # parse HTML and XML (sitemaps)
from dotenv import load_dotenv
from openai import AsyncOpenAI
from docx import Document  # python-docx: create Word documents
from docx.shared import Pt

load_dotenv()  # load .env so OPENAI_API_KEY and CRAWL_BASE_URL are available

BASE_URL = os.getenv("CRAWL_BASE_URL", "")
OUTPUT_DIR = Path("generated_docs")  # all DOCX/MD files go under this folder
# Browser-like headers so production (e.g. Render) is less likely to get 403/blocked
HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.9",
}

# URLs ending with these extensions are skipped (not HTML pages)
SKIP_EXTENSIONS = {
    ".xml", ".pdf", ".jpg", ".jpeg", ".png", ".gif", ".webp", ".svg", ".ico",
    ".css", ".js", ".woff", ".woff2", ".ttf", ".eot", ".mp4", ".webm", ".mp3",
    ".zip", ".rar", ".doc", ".docx", ".xls", ".xlsx", ".ppt", ".pptx",
    ".rss", ".atom", ".json", ".txt",
}
# Paths containing these patterns are skipped (e.g. WordPress assets, media)
SKIP_PATH_PATTERNS = re.compile(
    r"(/wp-content/|/wp-includes/|/static/|/assets/|/media/|/uploads/|/feed/|/tag/|/\d{4}/\d{2}/)",
    re.I
)

# Global job state: job_id -> { status, base_url, urls, groups, url_status, docs, ... }
# In production you would use a database or Redis instead of in-memory dict
jobs: dict = {}


# -------------------------
# URL normalization & filtering (accurate discovery)
# -------------------------
def normalize_url(url: str, base: str) -> str | None:
    """Normalize to absolute URL; return None if invalid or wrong domain."""
    if not url or not url.strip():
        return None
    url = url.strip().split("#")[0].strip()  # remove fragment (#section)
    if not url or url.startswith("mailto:") or url.startswith("tel:") or url.startswith("javascript:"):
        return None
    parsed = urlparse(url)
    base_parsed = urlparse(base)
    if parsed.scheme and parsed.scheme not in ("http", "https"):
        return None
    if parsed.netloc and parsed.netloc != base_parsed.netloc:  # must be same domain
        return None
    try:
        full = urljoin(base, url)  # make absolute (e.g. /about -> https://site.com/about)
    except Exception:
        return None
    p = urlparse(full)
    if p.scheme not in ("http", "https") or not p.netloc:
        return None
    if p.netloc != base_parsed.netloc:
        return None
    path = p.path or "/"
    path_lower = path.lower()
    if any(path_lower.endswith(ext) for ext in SKIP_EXTENSIONS):
        return None
    if SKIP_PATH_PATTERNS.search(path):
        return None
    # Normalize path: remove trailing slash (except for root), collapse ..
    path = path.rstrip("/") or "/"
    parts = path.split("/")
    resolved = []
    for part in parts:
        if part == "..":
            if resolved:
                resolved.pop()
        elif part and part != ".":
            resolved.append(part)
    path = "/" + "/".join(resolved) if resolved else "/"
    normalized = urlunparse((p.scheme, p.netloc, path, p.params, p.query or "", ""))
    return normalized


def url_to_safe_basename(url: str, max_len: int = 165) -> str:
    """Safe base name (no extension) for docx and .md files. No slashes so download URL works."""
    p = urlparse(url)
    path = (p.path or "/").strip("/") or "index"
    path = path.replace("/", "_")  # e.g. talon/amanda-walker -> talon_amanda-walker
    path = re.sub(r"[^\w\-.]", "_", path)  # only letters, digits, -, ., _
    path = re.sub(r"_+", "_", path).strip("_")
    if not path:
        path = "index"
    if len(path) > 140:
        path = path[:140]
    if p.query:
        qhash = hashlib.md5(p.query.encode()).hexdigest()[:8]
        path = f"{path}_{qhash}"
    if len(path) > max_len:
        path = path[: max_len - 10] + "_" + hashlib.md5(url.encode()).hexdigest()[:8]
    return path


def url_to_safe_filename(url: str, max_len: int = 180) -> str:
    """Turn URL into a safe .docx filename (single path component, no slashes)."""
    return url_to_safe_basename(url, max_len - 5) + ".docx"


def group_urls(urls: list[str]) -> dict[str, list[str]]:
    """Group URLs by first path segment (e.g. /talon/foo -> 'talon', /about -> 'about')."""
    groups: dict[str, list[str]] = {}
    for url in urls:
        p = urlparse(url)
        path = (p.path or "/").strip("/")
        if not path:
            seg = "index"
        else:
            seg = path.split("/")[0] or "index"
        if seg not in groups:
            groups[seg] = []
        groups[seg].append(url)
    return groups


# -------------------------
# Fetch
# -------------------------
# Discovery on prod (e.g. Render) can be slow or get blocked; long timeout + retry delay
DISCOVERY_FETCH_TIMEOUT = 90  # seconds per request
DISCOVERY_RETRY_DELAY = 3     # seconds to wait between retries (avoids rate limit)

async def fetch(session: aiohttp.ClientSession, url: str, timeout: int | float = 15) -> str:
    """Fetch URL with GET; return HTML as string, or empty string on error."""
    try:
        async with session.get(url, headers=HEADERS, timeout=timeout) as response:
            if response.status != 200:
                return ""
            return await response.text()
    except Exception:
        return ""


# -------------------------
# Sitemap discovery (accurate)
# -------------------------
async def _fetch_with_retry(session: aiohttp.ClientSession, url: str, retries: int = 3) -> str:
    """Fetch with retries and delay between; production often needs multiple tries."""
    for attempt in range(max(1, retries)):
        if attempt > 0:
            await asyncio.sleep(DISCOVERY_RETRY_DELAY)
        out = await fetch(session, url, timeout=DISCOVERY_FETCH_TIMEOUT)
        if out:
            return out
    return ""


async def find_sitemaps(session: aiohttp.ClientSession, base: str) -> list[str]:
    """Get sitemap URLs from robots.txt, or try common sitemap paths."""
    base = base.rstrip("/") + "/"
    robots_url = urljoin(base, "robots.txt")
    text = await _fetch_with_retry(session, robots_url)
    sitemaps = []
    for line in text.split("\n"):
        line = line.strip()
        if line.upper().startswith("SITEMAP:"):
            url = line.split(":", 1)[1].strip()
            if url:
                sitemaps.append(url)
    if not sitemaps:
        # Try common sitemap paths (sitemap.xml first); helps when robots.txt is blocked on prod
        for path in ("sitemap-index.xml", "sitemap_index.xml", "sitemap.xml"):
            sitemaps.append(urljoin(base, path))
    return sitemaps


def _extract_locs_from_xml(xml: str) -> list[str]:
    """Parse sitemap XML and return all <loc> URL strings."""
    soup = BeautifulSoup(xml, "xml")
    urls = []
    for loc in soup.find_all("loc"):
        if loc and loc.text:
            urls.append(loc.text.strip())
    return urls


async def get_all_sitemap_urls(session: aiohttp.ClientSession, base: str) -> list[str]:
    """Fetch sitemap(s), follow .xml index links, collect all leaf (page) URLs."""
    sitemaps_to_fetch = await find_sitemaps(session, base)
    seen_sitemaps: set[str] = set()
    leaf_urls: list[str] = []

    while sitemaps_to_fetch:
        url = sitemaps_to_fetch.pop()
        if url in seen_sitemaps:
            continue
        seen_sitemaps.add(url)
        xml = await _fetch_with_retry(session, url)
        if not xml or len(xml) < 10:
            continue
        locs = _extract_locs_from_xml(xml)
        for loc in locs:
            if not loc.startswith("http"):
                continue
            loc_lower = loc.lower()
            if loc_lower.endswith(".xml"):
                sitemaps_to_fetch.append(loc)  # another sitemap index
            else:
                leaf_urls.append(loc)  # actual page URL

    return leaf_urls


# -------------------------
# Fallback: discover links from base page (single-page/DFS-style)
# -------------------------
async def discover_urls_from_page(session: aiohttp.ClientSession, base: str, limit: int = 500) -> list[str]:
    """Fetch base page and extract same-domain links from that page only."""
    base_normalized = base.rstrip("/") + "/"
    fetch_url = base_normalized
    html = await _fetch_with_retry(session, fetch_url)
    links = extract_internal_links(html, base_normalized, fetch_url)
    seen: set[str] = set()
    out: list[str] = []
    root = normalize_url(fetch_url, base_normalized) or fetch_url
    if root not in seen:
        seen.add(root)
        out.append(root)
    for u in links:
        n = normalize_url(u, base_normalized)
        if n and n not in seen and len(out) < limit:
            seen.add(n)
            out.append(n)
    return out


# -------------------------
# Fallback: BFS discovery with normalized URLs
# -------------------------
async def discover_urls_bfs(session: aiohttp.ClientSession, base: str, limit: int = 1000) -> list[str]:
    """If no sitemap, discover URLs by crawling from base: fetch page, extract links, repeat (BFS)."""
    base_normalized = normalize_url(base, base) or base
    base_parsed = urlparse(base_normalized)
    seen: set[str] = set()
    queue: list[str] = [base_normalized]
    found: list[str] = []

    while queue and len(found) < limit:
        url = queue.pop(0)
        if url in seen:
            continue
        seen.add(url)
        found.append(url)
        html = await _fetch_with_retry(session, url)
        soup = BeautifulSoup(html, "lxml")
        for link in soup.find_all("a", href=True):
            href = link["href"].strip()
            normalized = normalize_url(href, base_normalized)
            if normalized and normalized not in seen:
                queue.append(normalized)

    return found


# -------------------------
# Extract clean text from HTML
# -------------------------
def html_to_text(html: str) -> str:
    """Remove scripts/nav/footer, get plain text from body. Max 100k chars."""
    if not html or len(html) < 50:
        return ""
    soup = BeautifulSoup(html, "lxml")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside", "form", "iframe"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return re.sub(r"\n{3,}", "\n\n", "\n".join(lines))[:1000000]


# -------------------------
# OpenAI/doc cleaning
# -------------------------
DOC_NOISE_LINE_PATTERN = re.compile(
    r"^(page\s*\d+|-\s*-\s*\d+\s+of\s+\d+\s*-\s*-|"
    r"\[?\s*learn\s+more\s*\]?|\[?\s*apply\s+now\s*\]?|apply\s+today|"
    r"read\s+more|click\s+here|submit|request\s+info|get\s+started|"
    r"next\s*page|previous\s*page|back\s+to\s+top|^\d+\s*$)$",
    re.I,
)


def _strip_doc_noise(text: str) -> str:
    """Remove page numbers, button text, and CTA noise from final document text."""
    if not text:
        return ""
    lines: list[str] = []
    for line in text.split("\n"):
        s = line.strip()
        if not s:
            lines.append("")
            continue
        if DOC_NOISE_LINE_PATTERN.search(s):
            continue
        s = re.sub(r"\s*\[\s*learn\s+more\s*\]\s*$", "", s, flags=re.I)
        s = re.sub(r"\s*\[\s*apply\s+now\s*\]\s*$", "", s, flags=re.I)
        s = re.sub(r"\s*apply\s+now\s*\.?\s*$", "", s, flags=re.I)
        s = re.sub(r"\s*learn\s+more\s*\.?\s*$", "", s, flags=re.I)
        if s.strip():
            lines.append(s)
    return re.sub(r"\n{3,}", "\n\n", "\n".join(lines)).strip()


def _normalize_readable_text(text: str) -> str:
    """Deterministic cleanup for stable, human-readable document text."""
    if not text:
        return ""
    lines = [ln.strip() for ln in text.splitlines()]
    out: list[str] = []
    seen_compact: set[str] = set()
    heading_alias = {
        "overview": "Overview",
        "program overview": "Program Overview",
        "description": "Program Description",
        "program description": "Program Description",
        "highlights": "Program Highlights",
        "learning outcomes": "Learning Outcomes",
        "outcomes": "Learning Outcomes",
        "requirements": "Requirements",
        "admissions": "Admissions",
        "fees": "Fees and Costs",
        "costs": "Fees and Costs",
        "faq": "FAQs",
        "faqs": "FAQs",
    }

    def _is_heading(s: str) -> bool:
        if not s or len(s) > 80 or s.endswith("."):
            return False
        if re.match(r"^[A-Z][A-Za-z0-9/&()\- ,]{2,}$", s):
            return True
        return s.lower().strip(":") in heading_alias

    for ln in lines:
        if not ln:
            out.append("")
            continue

        compact = re.sub(r"\s+", " ", ln).strip().lower()
        if compact in seen_compact and len(compact) > 20:
            continue
        seen_compact.add(compact)

        key = ln.lower().strip(":")
        if key in heading_alias:
            out.append(heading_alias[key])
            out.append("")
            continue

        if re.match(r"^[-*•]\s+", ln):
            ln = "• " + re.sub(r"^[-*•]\s+", "", ln).strip()
        elif re.match(r"^\d+\.\s+", ln):
            ln = "• " + re.sub(r"^\d+\.\s+", "", ln).strip()

        if ":" in ln and len(ln) < 140:
            k, v = ln.split(":", 1)
            if len(k.strip()) <= 40:
                ln = f"{k.strip()}: {v.strip()}"

        if _is_heading(ln):
            out.append(ln.strip(":"))
            out.append("")
            continue

        out.append(ln)

    text_out = "\n".join(out)
    text_out = re.sub(r"[ \t]+\n", "\n", text_out)
    return re.sub(r"\n{3,}", "\n\n", text_out).strip()


async def openai_format_content(client: AsyncOpenAI, raw_text: str) -> str:
    """Send raw text to OpenAI; return LLM-ready Markdown (with clear headings/lists)."""
    if not raw_text or len(raw_text) < 50:
        return raw_text
    chunk_size = 12000
    chunks = [raw_text[i : i + chunk_size] for i in range(0, len(raw_text), chunk_size)]
    formatted_parts = []
    prompt = """Format this web content into clean, highly readable Markdown for BOTH human DOCX and LLM use. Only format — do not add facts or change meaning.

ADAPTIVE STRUCTURE (important):
- Every page is different, so do NOT force a fixed template.
- Preserve the source meaning and natural section flow.
- Use headings/lists/tables only where they actually fit the source content.

OUTPUT RULES:
- Include short YAML front matter at top: `title`, `summary`, `source_url` (if available), `key_topics`.
- Use one clear `#` title, then `##`/`###` headings based on actual topic shifts.
- Keep key facts, numbers, dates, names, and URLs exactly accurate.
- Remove UI clutter/noise: page numbers, cookie/nav/footer text, [Learn More], Apply now, Read more, Click here.
- De-duplicate repeated lines/paragraphs.
- Keep paragraphs compact and clear.
- Use bullets for options/features/criteria.
- Use numbered lists for steps/processes.
- If tabular data exists, render as markdown table (`| Field | Value |`).
- If data is not tabular, do NOT force a table.
- Do NOT output HTML.

Content:
"""
    for chunk in chunks:
        try:
            r = await client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt + chunk}],
                temperature=0.0,
            )
            formatted_parts.append(r.choices[0].message.content or chunk)
        except Exception:
            formatted_parts.append(chunk)
    return "\n\n".join(formatted_parts)


# -------------------------
# One DOCX and one MD per URL (single page)
# -------------------------
def build_single_page_docx(url: str, plain_text: str, base_url: str) -> bytes:
    """Build DOCX from plain document text (no markdown). Returns file as bytes."""
    plain_text = _strip_doc_noise(plain_text)
    doc = Document()
    _style_doc_for_readability(doc)
    doc.add_heading(url, level=0)
    doc.add_paragraph(f"Program URL: {url}")
    doc.add_paragraph(f"Generated: {datetime.utcnow().isoformat()}Z")
    doc.add_paragraph("")
    _render_doc_blocks(doc, plain_text)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _markdown_to_doc_text(text: str) -> str:
    """Convert markdown/yaml-style content to plain document text for DOCX."""
    if not text:
        return ""
    out = []
    in_fence = False
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            out.append("")
            continue
        # Normalize smart quotes around fence markers from model outputs
        normalized = line.strip().strip('"').strip("'").strip("“”")
        lower_norm = normalized.lower()

        # Drop markdown/code fences (```yaml ... ```), including quoted variants
        if "```" in normalized or normalized.startswith("~~~"):
            in_fence = not in_fence
            continue
        if in_fence and lower_norm in ("yaml", "yml", "markdown", "md"):
            continue
        # Drop YAML frontmatter separators
        if normalized == "---":
            continue
        # Convert markdown table rows to readable key-value lines for DOCX text rendering.
        # Example: | Field | Value | -> Field: Value
        if normalized.startswith("|") and normalized.endswith("|"):
            cells = [c.strip() for c in normalized.strip("|").split("|")]
            if len(cells) >= 2:
                # Skip separator row like |---|---|
                if all(re.match(r"^:?-{2,}:?$", c) for c in cells):
                    continue
                key = cells[0]
                val = " | ".join(cells[1:]).strip()
                if val.lower() in {"null", "none", "n/a", "na", "unknown", "-"}:
                    continue
                if key and val:
                    out.append(f"{key}: {val}")
                    continue
        if re.match(r"^[-*_]{2,}\s*$", line):
            out.append("")
            continue
        # YAML key: value -> Human readable Key: value
        m_yaml = re.match(r"^([a-zA-Z0-9_ -]+):\s*(.*)$", normalized)
        if m_yaml:
            key = m_yaml.group(1).replace("_", " ").strip().title()
            val = m_yaml.group(2).strip()
            if val.lower() in {"null", "none", "n/a", "na", "unknown", "-"}:
                continue
            if val:
                out.append(f"{key}: {val}")
            else:
                # key_topics: -> Key Topics
                out.append(key)
            continue
        line = re.sub(r"^#+\s*", "", normalized).strip()
        line = line.replace("**", "").replace("__", "")
        line = re.sub(r"\*([^*]+)\*", r"\1", line)
        line = re.sub(r"_([^_]+)_", r"\1", line)
        if re.match(r"^[-*]\s+", line):
            line = "  • " + re.sub(r"^[-*]\s+", "", line)
        if re.match(r"^\d+\.\s+", line):
            line = re.sub(r"^\d+\.\s+", "  • ", line)
        line = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", line)
        line = re.sub(r"`([^`]+)`", r"\1", line)
        out.append(line)
    return re.sub(r"\n{3,}", "\n\n", "\n".join(out)).strip()


def _style_doc_for_readability(doc: Document) -> None:
    """Apply readable typography and spacing for humans."""
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        if name in doc.styles:
            doc.styles[name].font.name = "Calibri"
            doc.styles[name].font.size = Pt(size)


def _render_doc_blocks(doc: Document, text: str) -> None:
    """Render blocks with better heading/list/paragraph readability."""
    if not text:
        return
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if "\n" in block:
            for ln in [x.strip() for x in block.splitlines() if x.strip()]:
                if ln.startswith("•") or re.match(r"^[-*]\s+", ln):
                    txt = re.sub(r"^([•]|[-*])\s*", "", ln).strip()
                    try:
                        doc.add_paragraph(txt, style="List Bullet")
                    except Exception:
                        doc.add_paragraph(f"• {txt}")
                elif re.match(r"^\d+\.\s+", ln):
                    txt = re.sub(r"^\d+\.\s+", "", ln).strip()
                    try:
                        doc.add_paragraph(txt, style="List Number")
                    except Exception:
                        doc.add_paragraph(ln)
                else:
                    doc.add_paragraph(ln)
            continue

        if ":" in block and len(block) < 140:
            key, val = block.split(":", 1)
            if key and len(key.strip()) < 40:
                p = doc.add_paragraph()
                p.add_run(key.strip() + ": ").bold = True
                p.add_run(val.strip())
                continue

        if len(block) < 120 and not block.endswith("."):
            level = 1 if len(block) < 60 else 2
            doc.add_heading(block, level=level)
        else:
            doc.add_paragraph(block)


def build_single_page_md(url: str, markdown_text: str) -> str:
    """Build LLM-ready Markdown output. Keeps markdown structure and prepends metadata front matter."""
    body = _strip_doc_noise(markdown_text).strip()
    if not body.startswith("---"):
        meta = [
            "---",
            f'source_url: "{url}"',
            f'generated_at: "{datetime.utcnow().isoformat()}Z"',
            "format: llm_ready_markdown",
            "---",
            "",
        ]
        body = "\n".join(meta) + body
    return body


# -------------------------
# DFS: extract internal links from HTML (same domain, normalized, skip non-HTML)
# -------------------------
def extract_internal_links(html: str, base_url: str, current_url: str) -> list[str]:
    """Extract same-domain links from page HTML; normalize and filter."""
    if not html or len(html) < 50:
        return []
    soup = BeautifulSoup(html, "lxml")
    base_normalized = base_url.rstrip("/") + "/"
    out = []
    for a in soup.find_all("a", href=True):
        href = a["href"].strip()
        normalized = normalize_url(href, base_normalized)
        if normalized and normalized not in out:
            out.append(normalized)
    return out


# -------------------------
# Combined DOCX/MD for DFS (one doc per root with multiple sections)
# -------------------------
def build_combined_docx(root_url: str, sections: list[tuple[str, str]], base_url: str) -> bytes:
    """One DOCX: root URL as title, then Section per (url, plain_text)."""
    doc = Document()
    _style_doc_for_readability(doc)
    doc.add_heading(root_url, level=0)
    doc.add_paragraph(f"Source: {base_url.rstrip('/') or root_url}")
    doc.add_paragraph(f"Generated: {datetime.utcnow().isoformat()}Z")
    doc.add_paragraph("")
    for url, markdown_text in sections:
        plain_text = _normalize_readable_text(_strip_doc_noise(_markdown_to_doc_text(markdown_text)))
        doc.add_heading(f"Section: {url}", level=1)
        _render_doc_blocks(doc, plain_text)
        doc.add_paragraph("")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def build_combined_md(root_url: str, sections: list[tuple[str, str]]) -> str:
    """One LLM-ready MD: root metadata + markdown section per URL."""
    lines = [
        "---",
        f'source_root: "{root_url}"',
        f'generated_at: "{datetime.utcnow().isoformat()}Z"',
        "format: llm_ready_markdown",
        "---",
        "",
        "# Combined Crawl Document",
        "",
    ]
    for url, markdown_text in sections:
        lines.append(f"## Source URL: {url}")
        lines.append("")
        lines.append(_strip_doc_noise(markdown_text))
        lines.append("")
        lines.append("")
    return "\n".join(lines)


# -------------------------
# Process one URL: crawl → format → save DOCX + MD, update url_status
# -------------------------
def _set_url_failed(job_id: str, url: str, reason: str) -> None:
    """Mark a URL as failed and store reason so UI/API can show it."""
    job = jobs.get(job_id)
    if not job:
        return
    url_status = job.get("url_status") or {}
    url_status[url] = {"status": "failed", "docx": None, "md": None, "error": reason}
    job["url_status"] = url_status


async def process_one_url(
    session: aiohttp.ClientSession,
    openai_client: AsyncOpenAI,
    url: str,
    job_id: str,
    base_url: str,
) -> tuple[str | None, str | None]:
    """Fetch URL, extract text, format with OpenAI, convert to plain text, write DOCX and MD."""
    job = jobs.get(job_id)
    if not job:
        return None, None
    url_status = job.get("url_status") or {}
    url_status[url] = {"status": "crawling", "docx": None, "md": None}
    job["url_status"] = url_status

    # Use longer timeout + retry on prod so fetch from Render has time to get the page
    html = await _fetch_with_retry(session, url)
    raw_text = html_to_text(html)
    if not raw_text or len(raw_text) < 30:
        _set_url_failed(job_id, url, "Page returned no or too little content (check if site blocks this server)")
        return None, None

    try:
        formatted = await openai_format_content(openai_client, raw_text)
        plain = _normalize_readable_text(_strip_doc_noise(_markdown_to_doc_text(formatted)))
        basename = url_to_safe_basename(url)
        docx_name = basename + ".docx"
        md_name = basename + ".md"
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        out_dir = OUTPUT_DIR / job_id
        out_dir.mkdir(parents=True, exist_ok=True)

        (out_dir / docx_name).write_bytes(build_single_page_docx(url, plain, base_url))
        (out_dir / md_name).write_text(build_single_page_md(url, formatted), encoding="utf-8")

        job["docs"] = job.get("docs", []) + [docx_name, md_name]
        job["urls_done"] = job.get("urls_done", 0) + 1
        url_status[url] = {"status": "completed", "docx": docx_name, "md": md_name}
        job["url_status"] = url_status
        return docx_name, md_name
    except Exception as e:
        _set_url_failed(job_id, url, str(e))
        return None, None


def _ensure_url_status(job_id: str, urls: list[str]) -> None:
    """Ensure every URL in the list has an entry in job['url_status'] (pending if missing)."""
    job = jobs.get(job_id)
    if not job:
        return
    url_status = job.get("url_status") or {}
    for u in urls:
        if u not in url_status:
            url_status[u] = {"status": "pending", "docx": None, "md": None}
    job["url_status"] = url_status


def _path_normalized_for_similarity(path: str) -> str:
    """Normalize path for similarity check: no leading/trailing slash, collapsed."""
    path = (path or "/").strip("/") or ""
    parts = path.split("/")
    resolved = []
    for part in parts:
        if part == "..":
            if resolved:
                resolved.pop()
        elif part and part != ".":
            resolved.append(part)
    return "/".join(resolved)


def _paths_are_similar(path_a: str, path_b: str) -> bool:
    """
    True if both paths refer to the same page (equal or one path is a suffix of the other).
    E.g. /undergraduate-student-admissions and /admissions/undergraduate-student-admissions
    """
    a = _path_normalized_for_similarity(path_a)
    b = _path_normalized_for_similarity(path_b)
    if a == b:
        return True
    if not a or not b:
        return a == b
    # One path ends with the other as a path segment (e.g. b ends with a)
    return b == a or b.endswith("/" + a) or a.endswith("/" + b)


def _url_similar_to_any(candidate_url: str, sitemap_urls: set[str], base_netloc: str) -> bool:
    """True if candidate_url is exact match or path-similar to any sitemap URL (same domain)."""
    try:
        p = urlparse(candidate_url)
        if p.netloc and p.netloc != base_netloc:
            return False
        cand_path = (p.path or "/").rstrip("/") or "/"
        for s in sitemap_urls:
            sp = urlparse(s)
            if sp.netloc and sp.netloc != base_netloc:
                continue
            s_path = (sp.path or "/").rstrip("/") or "/"
            if _paths_are_similar(cand_path, s_path):
                return True
    except Exception:
        pass
    return False


# -------------------------
# DFS crawl from one root URL (depth 0 = root only; max_depth 1 = root + direct links; etc.)
# -------------------------
async def dfs_crawl_root(
    session: aiohttp.ClientSession,
    openai_client: AsyncOpenAI,
    root_url: str,
    base_url: str,
    max_depth: int,
    max_pages: int,
    max_links_per_page: int,
    skip_urls_in_sitemap: set[str] | None = None,
    job_id: str | None = None,
) -> list[tuple[str, int, str]]:
    """
    BFS traversal from root up to max_depth. Returns list of (url, depth, plain_text).
    If job_id is provided, updates job["dfs_progress"] after each page so frontend can show status.
    """
    base_url = base_url.rstrip("/") + "/"
    base_netloc = urlparse(base_url).netloc or urlparse(root_url).netloc
    sitemap_set = skip_urls_in_sitemap or set()
    visited: set[str] = set()
    queue: list[tuple[str, int]] = [(root_url, 0)]
    sections: list[tuple[str, int, str]] = []

    def _skip_sitemap_or_similar(link: str) -> bool:
        """True if link is in sitemap or path-similar to any sitemap URL (same page)."""
        if link in sitemap_set:
            return True
        return _url_similar_to_any(link, sitemap_set, base_netloc)

    def _update_progress(current_url: str) -> None:
        if job_id and job_id in jobs:
            jobs[job_id]["dfs_progress"] = {
                "root_url": root_url,
                "pages_crawled": len(sections),
                "total_limit": max_pages,
                "current_url": current_url,
            }

    print(f"DFS root: {root_url}")

    while queue and len(visited) < max_pages:
        url, depth = queue.pop(0)
        if url in visited:
            continue
        if depth > max_depth:
            continue
        visited.add(url)
        print(f"DFS visiting: depth={depth} {url}")

        html = await fetch(session, url)
        raw_text = html_to_text(html)
        if not raw_text or len(raw_text) < 30:
            continue
        formatted = await openai_format_content(openai_client, raw_text)
        sections.append((url, depth, formatted))
        _update_progress(url)

        if depth < max_depth and len(visited) < max_pages:
            links = extract_internal_links(html, base_url, url)
            # Skip links that are in sitemap or path-similar (e.g. /admissions/undergraduate... vs /undergraduate...)
            links_not_sitemap_or_similar = [u for u in links if not _skip_sitemap_or_similar(u)]
            new_links = [u for u in links_not_sitemap_or_similar if u not in visited][:max_links_per_page]
            for link in new_links:
                if link not in visited:
                    queue.append((link, depth + 1))

    if job_id and job_id in jobs:
        jobs[job_id].pop("dfs_progress", None)
    return sections


async def process_one_root_dfs(
    session: aiohttp.ClientSession,
    openai_client: AsyncOpenAI,
    root_url: str,
    job_id: str,
    base_url: str,
    max_depth: int,
    max_pages: int,
    max_links_per_page: int,
) -> tuple[str, str]:
    """Run DFS from root_url, build one combined DOCX + MD, save and update job. Returns (docx_name, md_name)."""
    job = jobs.get(job_id)
    if not job:
        return "", ""
    url_status = job.get("url_status") or {}
    url_status[root_url] = {"status": "crawling", "docx": None, "md": None}
    job["url_status"] = url_status

    job["dfs_progress"] = {"root_url": root_url, "pages_crawled": 0, "total_limit": max_pages, "current_url": ""}
    sitemap_urls = set(job.get("urls") or [])
    sections_with_depth = await dfs_crawl_root(
        session,
        openai_client,
        root_url,
        base_url,
        max_depth,
        max_pages,
        max_links_per_page,
        skip_urls_in_sitemap=sitemap_urls,
        job_id=job_id,
    )
    sections = [(url, plain) for url, _d, plain in sections_with_depth]
    if not sections:
        url_status[root_url] = {"status": "failed", "docx": None, "md": None, "error": "No content from root or linked pages"}
        job["url_status"] = url_status
        job.pop("dfs_progress", None)
        return "", ""

    basename = url_to_safe_basename(root_url) + "_dfs"
    docx_name = basename + ".docx"
    md_name = basename + ".md"
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    out_dir = OUTPUT_DIR / job_id
    out_dir.mkdir(parents=True, exist_ok=True)

    (out_dir / docx_name).write_bytes(build_combined_docx(root_url, sections, base_url))
    (out_dir / md_name).write_text(build_combined_md(root_url, sections), encoding="utf-8")

    job["docs"] = job.get("docs", []) + [docx_name, md_name]
    job["urls_done"] = job.get("urls_done", 0) + 1
    url_status[root_url] = {"status": "completed", "docx": docx_name, "md": md_name}
    job["url_status"] = url_status
    return docx_name, md_name


# -------------------------
# Discovery only: find URLs + group, do not crawl
# -------------------------
async def run_discovery_only(job_id: str, base_url: str) -> None:
    """Discover URLs (sitemap or BFS), normalize, group; set job status to 'discovered'. No crawling."""
    base_url = base_url.rstrip("/") + "/"
    jobs[job_id] = {
        "status": "discovering",
        "base_url": base_url,
        "total_urls": 0,
        "urls": [],
        "groups": {},
        "url_status": {},
        "urls_done": 0,
        "docs": [],
        "error": None,
    }
    try:
        # Warm-up: on prod (e.g. Render) discovery can finish in 1s with 1 URL if we start before network is ready
        await asyncio.sleep(8)
        timeout = aiohttp.ClientTimeout(total=DISCOVERY_FETCH_TIMEOUT, connect=30)
        async with aiohttp.ClientSession(timeout=timeout) as session:
            urls = await get_all_sitemap_urls(session, base_url)
            print(f"Discovery [{job_id}]: sitemap returned {len(urls)} URLs")
            base_no_slash = base_url.rstrip("/")
            only_base = not urls or (
                len(urls) == 1 and (normalize_url(urls[0], base_url) or "").rstrip("/") == base_no_slash
            )
            if only_base:
                urls = await discover_urls_from_page(session, base_url, limit=500)
                print(f"Discovery [{job_id}]: no sitemap/robots — DFS from page returned {len(urls)} URLs")
            elif not urls:
                urls = await discover_urls_bfs(session, base_url, limit=500)
                print(f"Discovery [{job_id}]: BFS fallback returned {len(urls)} URLs")
            seen = set()
            normalized_list = []
            for u in urls:
                n = normalize_url(u, base_url)
                if n and n not in seen:
                    seen.add(n)
                    normalized_list.append(n)
            groups = group_urls(normalized_list)
            url_status = {u: {"status": "pending", "docx": None, "md": None} for u in normalized_list}
            jobs[job_id]["status"] = "discovered"
            jobs[job_id]["urls"] = normalized_list
            jobs[job_id]["groups"] = groups
            jobs[job_id]["total_urls"] = len(normalized_list)
            jobs[job_id]["url_status"] = url_status
    except Exception as e:
        if job_id in jobs:
            jobs[job_id]["status"] = "failed"
            jobs[job_id]["error"] = str(e)


# -------------------------
# Crawl all selected URLs into one combined DOCX + MD
# -------------------------
async def run_crawl_urls_combined(job_id: str, urls: list[str], base_url: str) -> None:
    """Crawl selected URLs and generate one combined.docx + combined.md."""
    job = jobs.get(job_id)
    if not job or job.get("status") != "discovered":
        return
    _ensure_url_status(job_id, urls)
    job["status"] = "crawling"
    sections: list[tuple[str, str]] = []
    try:
        async with aiohttp.ClientSession() as session:
            openai_client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))
            for url in urls:
                url_status = job.get("url_status") or {}
                url_status[url] = {"status": "crawling", "docx": None, "md": None}
                job["url_status"] = url_status
                html = await _fetch_with_retry(session, url)
                raw_text = html_to_text(html)
                if not raw_text or len(raw_text) < 30:
                    _set_url_failed(job_id, url, "Page returned no or too little content")
                    continue
                try:
                    formatted = await openai_format_content(openai_client, raw_text)
                    sections.append((url, formatted))
                    job["urls_done"] = job.get("urls_done", 0) + 1
                    url_status[url] = {"status": "completed", "docx": "combined.docx", "md": "combined.md"}
                    job["url_status"] = url_status
                except Exception as e:
                    _set_url_failed(job_id, url, str(e))
        if sections:
            title = "Combined document — " + (base_url.rstrip("/") or "crawl")
            out_dir = OUTPUT_DIR / job_id
            out_dir.mkdir(parents=True, exist_ok=True)
            (out_dir / "combined.docx").write_bytes(build_combined_docx(title, sections, base_url))
            (out_dir / "combined.md").write_text(build_combined_md(title, sections), encoding="utf-8")
            job["docs"] = job.get("docs", []) + ["combined.docx", "combined.md"]
    except Exception as e:
        if job_id in jobs:
            jobs[job_id]["error"] = str(e)
    finally:
        if job_id in jobs:
            jobs[job_id]["status"] = "discovered"


# -------------------------
# Crawl a list of URLs: BFS (one doc per URL) or DFS (one combined doc per root)
# -------------------------
async def run_crawl_urls(
    job_id: str,
    urls: list[str],
    crawl_mode: str = "bfs",
    max_depth: int = 1,
    max_pages: int = 200,
    max_links_per_page: int = 20,
    combine_into_one_doc: bool = False,
) -> None:
    """
    BFS (default): crawl only the selected URLs, one DOCX + MD per URL (no link traversal).
    DFS: for each selected URL as root, crawl root + internal links up to max_depth; one combined DOCX + MD per root.
    """
    job = jobs.get(job_id)
    if not job or job.get("status") != "discovered":
        return
    base_url = job.get("base_url", "")
    _ensure_url_status(job_id, urls)
    job["status"] = "crawling"
    try:
        if combine_into_one_doc and crawl_mode == "bfs":
            await run_crawl_urls_combined(job_id, urls, base_url)
            return
        async with aiohttp.ClientSession() as session:
            openai_client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))
            if crawl_mode == "dfs":
                for root_url in urls:
                    await process_one_root_dfs(
                        session,
                        openai_client,
                        root_url,
                        job_id,
                        base_url,
                        max_depth=max_depth,
                        max_pages=max_pages,
                        max_links_per_page=max_links_per_page,
                    )
            else:
                # BFS: existing behavior — one doc per URL, no link traversal
                for url in urls:
                    await process_one_url(session, openai_client, url, job_id, base_url)
    except Exception as e:
        if job_id in jobs:
            jobs[job_id]["error"] = str(e)
    finally:
        if job_id in jobs:
            jobs[job_id].pop("dfs_progress", None)
            jobs[job_id]["status"] = "discovered"


# -------------------------
# Full pipeline: discover then crawl all one by one (legacy / "crawl all")
# -------------------------
async def run_crawl_pipeline(job_id: str, base_url: str) -> None:
    """Discover URLs, then crawl every URL one by one. Used when user clicks 'Crawl all'."""
    base_url = base_url.rstrip("/") + "/"
    jobs[job_id] = {
        "status": "discovering",
        "base_url": base_url,
        "total_urls": 0,
        "urls": [],
        "groups": {},
        "url_status": {},
        "urls_done": 0,
        "docs": [],
        "error": None,
    }

    try:
        async with aiohttp.ClientSession() as session:
            jobs[job_id]["status"] = "discovering"
            urls = await get_all_sitemap_urls(session, base_url)
            base_no_slash = base_url.rstrip("/")
            only_base = not urls or (
                len(urls) == 1 and (normalize_url(urls[0], base_url) or "").rstrip("/") == base_no_slash
            )
            if only_base:
                urls = await discover_urls_from_page(session, base_url, limit=500)
            elif not urls:
                urls = await discover_urls_bfs(session, base_url, limit=500)
            seen = set()
            normalized_list = []
            for u in urls:
                n = normalize_url(u, base_url)
                if n and n not in seen:
                    seen.add(n)
                    normalized_list.append(n)
            groups = group_urls(normalized_list)
            url_status = {u: {"status": "pending", "docx": None, "md": None} for u in normalized_list}
            jobs[job_id]["urls"] = normalized_list
            jobs[job_id]["groups"] = groups
            jobs[job_id]["total_urls"] = len(normalized_list)
            jobs[job_id]["url_status"] = url_status
            jobs[job_id]["status"] = "crawling"

            openai_client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))
            for url in normalized_list:
                await process_one_url(session, openai_client, url, job_id, base_url)

        if job_id in jobs:
            jobs[job_id]["status"] = "completed"
    except Exception as e:
        if job_id in jobs:
            jobs[job_id]["status"] = "failed"
            jobs[job_id]["error"] = str(e)


# -------------------------
# Standalone: save URLs only
# -------------------------
async def main_save_urls_only() -> None:
    """CLI helper: discover URLs and save to urls_fast.json (no crawling)."""
    base = BASE_URL.rstrip("/") + "/"
    async with aiohttp.ClientSession() as session:
        urls = await get_all_sitemap_urls(session, base)
        base_no_slash = base.rstrip("/")
        only_base = not urls or (
            len(urls) == 1 and (normalize_url(urls[0], base) or "").rstrip("/") == base_no_slash
        )
        if only_base:
            urls = await discover_urls_from_page(session, base, limit=500)
        elif not urls:
            urls = await discover_urls_bfs(session, base, limit=500)
        seen = set()
        normalized = []
        for u in urls:
            n = normalize_url(u, base)
            if n and n not in seen:
                seen.add(n)
                normalized.append(n)
        data = {"website": base, "total_urls": len(normalized), "urls": sorted(normalized)}
        with open("urls_fast.json", "w") as f:
            json.dump(data, f, indent=2)
        print("Saved to urls_fast.json", len(normalized), "URLs")


if __name__ == "__main__":
    job_id = str(uuid.uuid4())[:8]
    asyncio.run(run_crawl_pipeline(job_id, BASE_URL))
    print("Job", job_id, jobs.get(job_id))
